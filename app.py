from flask import Flask, render_template, request, redirect, url_for, session, send_file
import pandas as pd
import os
from docx import Document
import io
import random
import docx2txt
import PyPDF2
import difflib
import openai
import traceback
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
import torch
import google.generativeai as genai
from datetime import datetime
import re
from textwrap import dedent

from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = "secret123"

# Allowed file extensions 
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Dummy user DB
users = {
    "nidasiddiqui": {
        "password": "1234",
        "email": "nidasiddiqui2720@gmail.com",
        "role": "AI Career Explorer"
    }
}

# ---------- Helper Functions ----------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path):
    return docx2txt.process(file_path)

# ---------- Load your CSV globally ----------
df = pd.read_csv(r"database.csv")
df.columns = df.columns.str.strip().str.lower()

# Create mappings
career_skill_map = df.groupby("suggested_career")["skills"].apply(list).to_dict()
career_industry_map = df.set_index("suggested_career")["major_application"].to_dict()

# ---------- Career Suggestion Logic ----------
def map_user_skills(user_skills):
    all_skills = df["skills"].unique().tolist()
    mapped = []
    for skill in user_skills:
        match = difflib.get_close_matches(skill.lower(), [s.lower() for s in all_skills], n=1, cutoff=0.6)
        if match:
            for s in all_skills:
                if s.lower() == match[0]:
                    mapped.append(s)
    return mapped

def suggest_career(user_skills):
    user_skills = map_user_skills(user_skills)
    results = []

    for career, required_skills in career_skill_map.items():
        required_skills_set = set(s.lower() for s in required_skills)
        user_skills_set = set(s.lower() for s in user_skills)

        matched = required_skills_set & user_skills_set
        missing = required_skills_set - user_skills_set
        match_percent = (len(matched) / len(required_skills_set) * 100) if required_skills_set else 0

        results.append({
            "career": career,
            "industry": career_industry_map.get(career, "Unknown"),
            "match_percent": round(match_percent, 1),
            "required": list(required_skills_set),
            "matched": list(matched),
            "missing": list(missing)
        })

    results = sorted(results, key=lambda x: x["match_percent"], reverse=True)
    return results[:5]  # top 5 matches


# ---------- Routes ----------
@app.route('/')
def home():
    user = session.get("user")
    return render_template("index.html", user=user)

@app.route('/about')
def about():
    user = session.get("user")
    return render_template("about.html", user=user)

@app.route('/features')
def features():
    user = session.get("user")
    return render_template("features.html", user=user)

@app.route('/contact')
def contact():
    user = session.get("user")
    return render_template("contact.html", user=user)

@app.route('/demo')
def demo():
    return render_template('demo.html')  # create a demo.html template

@app.route('/image')
def image():
    return render_template('image.html')

# ---------- LLM Configuration (OpenAI, Gemini, Local) ----------

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "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").strip()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyDlHLNl29HjDKBIW7fHTY9K-i4SK3gp2mo").strip()

if OPENAI_API_KEY:
    openai.api_key = OPENAI_API_KEY

gemini_model = None
if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        # Use a lightweight, cheap model
        gemini_model = genai.GenerativeModel("gemini-1.5-flash")
    except Exception as e:
        print("Gemini configuration failed:", e)
        gemini_model = None

USE_LOCAL = True  # Local model is free and runs without an external API
local_model = None
local_tokenizer = None
try:
    if USE_LOCAL:
        model_name = "google/flan-t5-small"   # lightweight
        local_tokenizer = AutoTokenizer.from_pretrained(model_name)
        local_model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
except Exception as e:
    # if transformers not installed or model download fails, ignore
    local_model = None
    local_tokenizer = None
    print("Local model unavailable:", e)

def generate_with_openai(subject, num=10):
    prompt = f"Generate {num} interview questions and short (1-2 sentence) answers for the subject: {subject}. Format: Q1: <question> A1: <answer>"
    if not OPENAI_API_KEY:
        raise RuntimeError("OpenAI API key not configured")
    resp = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1500,
        temperature=0.7,
    )
    return resp["choices"][0]["message"]["content"]


def generate_with_gemini(subject, num=10):
    if gemini_model is None:
        raise RuntimeError("Gemini model not configured")
    prompt = (
        f"Generate {num} interview questions and short (1-2 sentence) answers for the "
        f"subject: {subject}. Format strictly as lines starting with Q1:, A1:, Q2:, A2:, etc."
    )
    response = gemini_model.generate_content(prompt)
    return response.text or ""

def generate_with_local(subject, num=10):
    if local_model is None:
        raise RuntimeError("Local model not available")
    prompt = f"Generate {num} interview questions and short (1-2 sentence) answers for the subject: {subject}. Format: Q1: <question> A1: <answer>"
    inputs = local_tokenizer(prompt, return_tensors="pt")
    out = local_model.generate(**inputs, max_new_tokens=512)
    return local_tokenizer.decode(out[0], skip_special_tokens=True)

QUESTION_PREFIX = re.compile(r"^(?:q(?:uestion)?\s*\d*|[0-9]+)(?:[:.)\-\]]\s*|\s+)", re.IGNORECASE)
ANSWER_PREFIX = re.compile(r"^(?:a(?:nswer)?\s*\d*)(?:[:.)\-\]]\s*|\s+)", re.IGNORECASE)


def _strip_prefix(line: str, pattern: re.Pattern):
    line = line.strip()
    if not line:
        return None
    match = pattern.match(line)
    if match:
        return line[match.end():].strip() or None
    return None


def parse_qa(text):
    """Parse Q&A pairs from raw LLM output with multiple heuristics."""
    lines = [l.rstrip() for l in text.splitlines() if l.strip()]
    qa = []
    q, a = "", ""

    for line in lines:
        question = _strip_prefix(line, QUESTION_PREFIX)
        answer = _strip_prefix(line, ANSWER_PREFIX)

        if question is not None:
            if q:
                qa.append((q.strip(), a.strip()))
                a = ""
            q = question or line.strip()
            continue

        if answer is not None:
            a = (a + " " + answer).strip()
            continue

        if q:
            a = (a + " " + line.strip()).strip()

    if q:
        qa.append((q.strip(), a.strip()))

    if qa:
        return qa

    # Fallback parsing: split blocks by blank lines, assume first line question
    blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
    for block in blocks:
        parts = [p.strip() for p in block.split("\n") if p.strip()]
        if len(parts) >= 2:
            qa.append((parts[0], " ".join(parts[1:])))

    return qa


def parse_project_ideas(text, limit):
    candidates = []
    for line in text.splitlines():
        clean = line.strip()
        if not clean:
            continue
        clean = clean.lstrip("-•0123456789. )(").strip()
        if clean:
            candidates.append(clean)

    if not candidates:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for para in paragraphs:
            clean = para.strip()
            if clean:
                candidates.append(clean)

    seen = set()
    ideas = []
    for idea in candidates:
        normalized = idea.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        ideas.append(idea)
        if len(ideas) >= limit:
            break
    return ideas


def generate_ai_project_ideas(topic, count=5):
    if not OPENAI_API_KEY:
        raise RuntimeError("OpenAI API key not configured for project ideas")

    prompt = dedent(
        f"""
        You are a career mentor. Generate {count} distinct, practical project ideas for the topic "{topic}".
        Requirements:
        - Target intermediate learners
        - Each idea must be one sentence max 30 words
        - Start each idea on a new line prefixed with "- "
        - Mention tools/tech stack briefly
        """
    ).strip()

    resp = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=600,
        temperature=0.7,
    )
    text = resp["choices"][0]["message"]["content"]
    ideas = parse_project_ideas(text, count)
    return ideas


SUBJECT_TEMPLATES = {
    "python": [
        ("What are Python's key features?", "Discuss readability, dynamic typing, the extensive standard library, and multi-paradigm support."),
        ("Explain Python's memory management model.", "Describe reference counting, garbage collection, and the role of the private heap."),
        ("Difference between list, tuple, and set in Python?", "Lists are mutable sequences, tuples are immutable, and sets are unordered collections of unique items."),
        ("How do decorators work in Python?", "They wrap callables to modify behavior, typically defined as functions returning inner functions."),
        ("What is the GIL and why does it matter?", "The Global Interpreter Lock allows only one Python bytecode thread at a time, limiting CPU-bound multi-threading."),
    ],
    "machine learning": [
        ("What is the bias-variance trade-off?", "It balances underfitting (bias) and overfitting (variance) when selecting model complexity."),
        ("Explain supervised vs unsupervised learning.", "Supervised learning uses labeled data; unsupervised finds structure in unlabeled data."),
        ("What is overfitting and how to avoid it?", "Overfitting memorizes noise; mitigate with regularization, cross-validation, and more data."),
        ("Describe gradient descent.", "An iterative optimization algorithm that updates parameters opposite the gradient of the loss."),
        ("When would you use a confusion matrix?", "To analyze classification performance across true/false positives and negatives."),
    ],
    "data structures": [
        ("Compare arrays and linked lists.", "Arrays offer O(1) indexing but costly resizing; linked lists allow O(1) insert/delete but O(n) indexing."),
        ("What is a hash table?", "A structure mapping keys to values via hash functions, offering average O(1) lookups."),
        ("Describe a balanced binary search tree.", "BST variants (AVL, Red-Black) keep height O(log n) to guarantee logarithmic operations."),
        ("Explain stack vs queue use cases.", "Stacks (LIFO) suit recursion/undo; queues (FIFO) suit scheduling and BFS."),
        ("When would you use a heap?", "When you need quick access to the smallest/largest element, e.g., priority queues."),
    ],
    "web development": [
        ("Explain the client-server model in web apps.", "Clients send HTTP requests; servers process logic and return resources or data."),
        ("What is REST and why is it popular?", "A stateless architectural style leveraging HTTP verbs and uniform resource identifiers."),
        ("How do CSS flexbox and grid differ?", "Flexbox manages one-dimensional layouts; CSS Grid handles two-dimensional structures."),
        ("Explain the critical rendering path.", "Steps the browser performs to convert HTML/CSS/JS into pixels; optimizing reduces blocking."),
        ("What role do build tools like Webpack play?", "They bundle assets, transpile modern syntax, and optimize delivery for production."),
    ],
}

GENERIC_TEMPLATES = [
    ("Explain concept #{idx} in {subject}.", "Provide a concise overview with practical context for concept #{idx}."),
    ("How is {subject} applied in real projects #{idx}?", "Describe an example scenario showing where it fits best."),
    ("What challenges arise with {subject} #{idx}?", "Highlight trade-offs, limitations, or pitfalls to watch."),
    ("Outline best practices for {subject} #{idx}.", "List actionable guidelines or architectural considerations professionals follow."),
    ("How does {subject} compare to alternatives #{idx}?", "Contrast strengths, weaknesses, and selection criteria."),
]


def generate_template_questions(subject: str, num: int):
    key = (subject or "").strip().lower()
    subject_label = subject or "this topic"
    templates = SUBJECT_TEMPLATES.get(key)

    if not templates:
        templates = GENERIC_TEMPLATES

    qa = []
    for idx in range(num):
        template_q, template_a = templates[idx % len(templates)]
        qa.append(
            (
                template_q.replace("{subject}", subject_label).replace("#{idx}", str(idx + 1)),
                template_a.replace("{subject}", subject_label).replace("#{idx}", str(idx + 1)),
            )
        )
    return qa

@app.route("/interview")
def interview():
    return render_template("interview.html")

@app.route("/interview_result", methods=["POST"])
def interview_result():
    subject = request.form.get("subject","").strip()
    requested_n = int(request.form.get("num_questions", 10))   # allow selecting count
    provider = request.form.get("provider", "openai").lower()

    # Chunk size to avoid huge single call
    chunk = 10
    generated_text = ""
    try:
        parts = []
        # Decide which provider to use
        if provider == "openai":
            for _ in range((requested_n + chunk - 1) // chunk):
                n = min(chunk, requested_n - len(parts) * chunk)
                parts.append(generate_with_openai(subject, n))
        elif provider == "gemini":
            for _ in range((requested_n + chunk - 1) // chunk):
                n = min(chunk, requested_n - len(parts) * chunk)
                parts.append(generate_with_gemini(subject, n))
        elif provider == "local":
            for _ in range((requested_n + chunk - 1) // chunk):
                n = min(chunk, requested_n - len(parts) * chunk)
                parts.append(generate_with_local(subject, n))
        else:
            raise RuntimeError(f"Unknown provider: {provider}")
        generated_text = "\n".join(parts)
    except Exception as e:
        print(f"{provider} generation failed:", e)
        traceback.print_exc()
        # fallback to local model if available
        try:
            parts = []
            for _ in range((requested_n + chunk - 1) // chunk):
                n = min(chunk, requested_n - len(parts) * chunk)
                parts.append(generate_with_local(subject, n))
            generated_text = "\n".join(parts)
            provider = "local"
        except Exception as e2:
            print("Local generation failed:", e2)
            # final fallback: small static generator
            generated_text = ""
            for i in range(1, requested_n + 1):
                generated_text += (
                    f"Q{i}: What is {subject} concept {i}?\n"
                    f"A{i}: Short answer for concept {i}.\n\n"
                )
            provider = "static"

    qa_list = parse_qa(generated_text)

    unique = []
    seen = set()
    for q_text, a_text in qa_list:
        norm = (q_text.strip().lower(), a_text.strip().lower())
        if norm in seen:
            continue
        seen.add(norm)
        unique.append((q_text.strip(), a_text.strip()))

    qa_list = unique

    if not qa_list:
        qa_list = generate_template_questions(subject, requested_n)
    elif len(qa_list) < requested_n:
        supplement = generate_template_questions(subject, requested_n - len(qa_list))
        qa_list.extend(supplement)

    qa_list = qa_list[:requested_n]
    generated_at = datetime.utcnow().strftime("%d %b %Y, %H:%M UTC")
    return render_template(
        "interview_result.html",
        subject=subject,
        questions=qa_list,
        provider=provider,
        total_questions=len(qa_list),
        num_requested=requested_n,
        generated_at=generated_at,
    )

@app.route("/resume_form", methods=["GET", "POST"])
def resume_form():
    if request.method == "POST":
        # Collect form data
        resume_data = {
            "name": request.form.get("name"),
            "email": request.form.get("email"),
            "phone": request.form.get("phone"),
            "address": request.form.get("address"),
            "education": request.form.get("education"),
            "experience": request.form.get("experience"),
            "skills": request.form.get("skills"),
            "projects": request.form.get("projects"),
            "achievements": request.form.get("achievements"),
        }
        # Render preview page with resume_data
        return render_template("resume_preview.html", resume=resume_data)

    return render_template("resume_form.html")


@app.route("/download_resume", methods=["POST"])
def download_resume():
    # Collect data from hidden inputs
    resume_data = {
        "name": request.form.get("name"),
        "email": request.form.get("email"),
        "phone": request.form.get("phone"),
        "address": request.form.get("address"),
        "education": request.form.get("education"),
        "experience": request.form.get("experience"),
        "skills": request.form.get("skills"),
        "projects": request.form.get("projects"),
        "achievements": request.form.get("achievements"),
    }

    # Create Word document
    doc = Document()
    doc.add_heading(resume_data["name"], 0)
    doc.add_paragraph(f"Email: {resume_data['email']} | Phone: {resume_data['phone']}")
    doc.add_paragraph(f"Address: {resume_data['address']}")
    doc.add_heading("Education", level=1)
    doc.add_paragraph(resume_data["education"])
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(resume_data["experience"])
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(resume_data["skills"])
    doc.add_heading("Projects", level=1)
    doc.add_paragraph(resume_data["projects"])
    doc.add_heading("Achievements", level=1)
    doc.add_paragraph(resume_data["achievements"])

    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"{resume_data['name']}_Resume.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/career_cv", methods=["GET", "POST"])
def career_cv():
    if request.method == "POST":
        # Get user details from form
        name = request.form.get("name")
        email = request.form.get("email")
        phone = request.form.get("phone")
        education = request.form.get("education")
        skills = request.form.get("skills")
        experience = request.form.get("experience")
        achievements = request.form.get("achievements")
        
        # Combine into a dict to send to preview page
        user_data = {
            "name": name,
            "email": email,
            "phone": phone,
            "education": education,
            "skills": skills,
            "experience": experience,
            "achievements": achievements
        }

        return render_template("career_preview.html", data=user_data)

    return render_template("career.html")

project_suggestions = {
    "python": [
        "Build a web scraper using BeautifulSoup",
        "Create a personal expense tracker with Tkinter",
        "Develop a mini chatbot using Python and NLP",
        "Automate email reports using Python scripts"
    ],
    "web development": [
        "Build a responsive portfolio website",
        "Create a task manager app using React + Firebase",
        "Develop a blog platform with Flask/Django",
        "Create a real-time chat app with Socket.io"
    ],
    "data science": [
        "Analyze global COVID-19 dataset and visualize trends",
        "Build a predictive model for house prices",
        "Create a recommendation system for movies",
        "Perform sentiment analysis on Twitter data"
    ]
}

# -------- Routes --------
@app.route("/project_form", methods=["GET", "POST"])
def project_form():
    if request.method == "POST":
        topic_raw = request.form.get("topic", "").strip()
        count = request.form.get("num_projects", "4")
        try:
            count = max(3, min(8, int(count)))
        except ValueError:
            count = 4

        suggestions = []
        provider_used = "openai"
        error_message = None

        try:
            suggestions = generate_ai_project_ideas(topic_raw, count)
        except Exception as ai_err:
            provider_used = "fallback"
            error_message = f"AI suggestion failed: {ai_err}"
            topic_key = topic_raw.lower()
            fallback_pool = project_suggestions.get(
                topic_key,
                [
                    f"Design a mini research project exploring {topic_raw or 'your topic'}.",
                    f"Build a knowledge blog documenting {topic_raw or 'key'} concepts.",
                    f"Create a tutorial series showing how to apply {topic_raw or 'this skill'} to a real dataset.",
                ],
            )

            if len(fallback_pool) >= count:
                suggestions = random.sample(fallback_pool, count)
            else:
                suggestions = list(fallback_pool)
                idx = 1
                while len(suggestions) < count:
                    suggestions.append(
                        f"Prototype idea #{idx} demonstrating {topic_raw or 'this skill'} in a small real-world scenario."
                    )
                    idx += 1

        topic_display = topic_raw or "Your Topic"
        return render_template(
            "project_result.html",
            topic=topic_display,
            suggestions=suggestions,
            provider=provider_used,
            total=count,
            error_message=error_message,
        )
    return render_template("project_form.html")

industry_data = [
    {"skill": "Python", "current_demand": "High", "future_demand": "Very High"},
    {"skill": "Java", "current_demand": "High", "future_demand": "High"},
    {"skill": "C++", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "C", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "JavaScript", "current_demand": "High", "future_demand": "High"},
    {"skill": "HTML", "current_demand": "High", "future_demand": "High"},
    {"skill": "CSS", "current_demand": "High", "future_demand": "High"},
    {"skill": "SQL", "current_demand": "High", "future_demand": "High"},
    {"skill": "R", "current_demand": "Moderate", "future_demand": "High"},
    {"skill": "PHP", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "Ruby", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "Kotlin", "current_demand": "Moderate", "future_demand": "High"},
    {"skill": "Swift", "current_demand": "High", "future_demand": "High"},
    {"skill": "TypeScript", "current_demand": "High", "future_demand": "High"},
    {"skill": "Go (Golang)", "current_demand": "Moderate", "future_demand": "High"},
    {"skill": "MATLAB", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "Excel", "current_demand": "High", "future_demand": "High"},
    {"skill": "PowerBI", "current_demand": "High", "future_demand": "High"},
    {"skill": "Tableau", "current_demand": "High", "future_demand": "High"},
    {"skill": "Git/GitHub", "current_demand": "High", "future_demand": "High"},
    {"skill": "Linux", "current_demand": "High", "future_demand": "High"},
    {"skill": "Docker", "current_demand": "High", "future_demand": "Very High"},
    {"skill": "Kubernetes", "current_demand": "Moderate", "future_demand": "High"},
    {"skill": "React", "current_demand": "High", "future_demand": "High"},
    {"skill": "Angular", "current_demand": "Moderate", "future_demand": "High"},
    {"skill": "Vue.js", "current_demand": "Moderate", "future_demand": "High"},
    {"skill": "Node.js", "current_demand": "High", "future_demand": "High"},
    {"skill": "Flutter", "current_demand": "High", "future_demand": "High"},
    {"skill": "Dart", "current_demand": "Moderate", "future_demand": "High"},
    {"skill": "React Native", "current_demand": "High", "future_demand": "High"},
    {"skill": "Android Development", "current_demand": "High", "future_demand": "High"},
    {"skill": "iOS Development", "current_demand": "High", "future_demand": "High"},
    {"skill": "WordPress", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "SEO", "current_demand": "High", "future_demand": "High"},
    {"skill": "Digital Marketing", "current_demand": "High", "future_demand": "High"},
    {"skill": "Content Writing", "current_demand": "High", "future_demand": "High"},
    {"skill": "Adobe Photoshop", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "Adobe Illustrator", "current_demand": "Moderate", "future_demand": "Moderate"},
    {"skill": "Canva", "current_demand": "High", "future_demand": "High"},
    {"skill": "Communication Skills", "current_demand": "High", "future_demand": "High"},
    {"skill": "Teamwork", "current_demand": "High", "future_demand": "High"},
    {"skill": "Problem Solving", "current_demand": "High", "future_demand": "High"},
    {"skill": "Critical Thinking", "current_demand": "High", "future_demand": "High"},
    {"skill": "Time Management", "current_demand": "High", "future_demand": "High"},
    {"skill": "Project Management", "current_demand": "High", "future_demand": "High"},
     {"skill": "Artificial Intelligence", "current_demand": "", "future_demand": ""},
    {"skill": "Machine Learning", "current_demand": "", "future_demand": ""},
    {"skill": "Deep Learning", "current_demand": "", "future_demand": ""},
    {"skill": "Data Science", "current_demand": "", "future_demand": ""},
    {"skill": "Big Data Analytics", "current_demand": "", "future_demand": ""},
    {"skill": "Data Engineering", "current_demand": "", "future_demand": ""},
    {"skill": "Predictive Analytics", "current_demand": "", "future_demand": ""},
    {"skill": "Statistical Modeling", "current_demand": "", "future_demand": ""},
    {"skill": "Cloud Computing", "current_demand": "", "future_demand": ""},
    {"skill": "AWS", "current_demand": "", "future_demand": ""},
    {"skill": "Azure", "current_demand": "", "future_demand": ""},
    {"skill": "Google Cloud Platform", "current_demand": "", "future_demand": ""},
    {"skill": "DevOps", "current_demand": "", "future_demand": ""},
    {"skill": "Full Stack Development", "current_demand": "", "future_demand": ""},
    {"skill": "Backend Development", "current_demand": "", "future_demand": ""},
    {"skill": "Frontend Development", "current_demand": "", "future_demand": ""},
    {"skill": "Web Development (React/Angular/Vue)", "current_demand": "", "future_demand": ""},
    {"skill": "Mobile App Development", "current_demand": "", "future_demand": ""},
    {"skill": "Internet of Things (IoT)", "current_demand": "", "future_demand": ""},
    {"skill": "Edge Computing", "current_demand": "", "future_demand": ""},
    {"skill": "Cybersecurity", "current_demand": "", "future_demand": ""},
    {"skill": "Ethical Hacking", "current_demand": "", "future_demand": ""},
    {"skill": "Blockchain Development", "current_demand": "", "future_demand": ""},
    {"skill": "Smart Contract Development", "current_demand": "", "future_demand": ""},
    {"skill": "Quantum Computing Basics", "current_demand": "", "future_demand": ""},
    {"skill": "Robotic Process Automation (RPA)", "current_demand": "", "future_demand": ""},
    {"skill": "Natural Language Processing (NLP)", "current_demand": "", "future_demand": ""},
    {"skill": "Prompt Engineering", "current_demand": "", "future_demand": ""},
    {"skill": "Generative AI", "current_demand": "", "future_demand": ""},
    {"skill": "Virtual Reality (VR) / Augmented Reality (AR)", "current_demand": "", "future_demand": ""},
    {"skill": "UX/UI Design", "current_demand": "", "future_demand": ""},
    {"skill": "User Experience Research", "current_demand": "", "future_demand": ""},
    {"skill": "Graphic Design", "current_demand": "", "future_demand": ""},
    {"skill": "Motion Graphics & Animation", "current_demand": "", "future_demand": ""},
    {"skill": "Digital Marketing", "current_demand": "", "future_demand": ""},
    {"skill": "Search Engine Optimization (SEO)", "current_demand": "", "future_demand": ""},
    {"skill": "Social Media Marketing", "current_demand": "", "future_demand": ""},
    {"skill": "Content Strategy & Creation", "current_demand": "", "future_demand": ""},
    {"skill": "Growth Hacking", "current_demand": "", "future_demand": ""},
    {"skill": "Product Management", "current_demand": "", "future_demand": ""},
    {"skill": "Agile & Scrum Methods", "current_demand": "", "future_demand": ""},
    {"skill": "Project Management", "current_demand": "", "future_demand": ""},
    {"skill": "Design Thinking", "current_demand": "", "future_demand": ""},
    {"skill": "Critical Thinking & Problem Solving", "current_demand": "", "future_demand": ""},
    {"skill": "Analytical Thinking", "current_demand": "", "future_demand": ""},
    {"skill": "Creative Thinking", "current_demand": "", "future_demand": ""},
    {"skill": "Leadership & Social Influence", "current_demand": "", "future_demand": ""},
    {"skill": "Emotional Intelligence", "current_demand": "", "future_demand": ""},
    {"skill": "Adaptability & Flexibility", "current_demand": "", "future_demand": ""},
    {"skill": "Communication Skills (Written & Verbal)", "current_demand": "", "future_demand": ""},
    {"skill": "Collaboration & Teamwork", "current_demand": "", "future_demand": ""},
    {"skill": "Resilience & Stress Management", "current_demand": "", "future_demand": ""},
    {"skill": "Technological Literacy", "current_demand": "", "future_demand": ""},
    {"skill": "Data Visualization (Tableau, PowerBI)", "current_demand": "", "future_demand": ""},
    {"skill": "SQL & Database Management", "current_demand": "", "future_demand": ""},
    {"skill": "NoSQL Databases (MongoDB, Cassandra)", "current_demand": "", "future_demand": ""},
    {"skill": "Event‑Driven Architecture", "current_demand": "", "future_demand": ""},
    {"skill": "APIs & Microservices", "current_demand": "", "future_demand": ""},
    {"skill": "Containerization (Docker & Kubernetes)", "current_demand": "", "future_demand": ""},
    {"skill": "Serverless Architecture", "current_demand": "", "future_demand": ""},
    {"skill": "QA & Test Automation", "current_demand": "", "future_demand": ""},
    {"skill": "Behavior‑Driven Development", "current_demand": "", "future_demand": ""},
    {"skill": "Software Architecture", "current_demand": "", "future_demand": ""},
    {"skill": "Systems Thinking", "current_demand": "", "future_demand": ""},
    {"skill": "Business Intelligence", "current_demand": "", "future_demand": ""},
    {"skill": "Supply Chain Analytics", "current_demand": "", "future_demand": ""},
    {"skill": "Sustainability / Green Skills", "current_demand": "", "future_demand": ""},
    {"skill": "Environmental Science & Policy", "current_demand": "", "future_demand": ""},
    {"skill": "Renewable Energy Technologies", "current_demand": "", "future_demand": ""},
    {"skill": "Healthcare Informatics", "current_demand": "", "future_demand": ""},
    {"skill": "Bioinformatics", "current_demand": "", "future_demand": ""},
    {"skill": "Genetic Data Analysis", "current_demand": "", "future_demand": ""},
    {"skill": "Telemedicine Technologies", "current_demand": "", "future_demand": ""},
    {"skill": "Remote & Hybrid Work Tools", "current_demand": "", "future_demand": ""},
    {"skill": "Virtual Collaboration Tools", "current_demand": "", "future_demand": ""},
    {"skill": "Coaching & Mentoring Skills", "current_demand": "", "future_demand": ""},
    {"skill": "Instructional Design & E‑Learning", "current_demand": "", "future_demand": ""},
    {"skill": "Change Management", "current_demand": "", "future_demand": ""},
    {"skill": "Talent Development & Reskilling", "current_demand": "", "future_demand": ""},
    {"skill": "Customer Experience (CX) Strategy", "current_demand": "", "future_demand": ""},
    {"skill": "Sales Automation & CRM", "current_demand": "", "future_demand": ""},
    {"skill": "FinTech Skills (Crypto, Digital Payments)", "current_demand": "", "future_demand": ""},
    {"skill": "Trading & Quantitative Finance", "current_demand": "", "future_demand": ""},
    {"skill": "Risk Management & Compliance", "current_demand": "", "future_demand": ""},
    {"skill": "Business Strategy & Policy Analysis", "current_demand": "", "future_demand": ""},
    {"skill": "Market Research & Trend Forecasting", "current_demand": "", "future_demand": ""},
    {"skill": "Voice User Interface (VUI) Design", "current_demand": "", "future_demand": ""},
    {"skill": "Video Editing & Short‑form Content", "current_demand": "", "future_demand": ""},
    {"skill": "Podcast Production & Audio Editing", "current_demand": "", "future_demand": ""},
    {"skill": "Translation & Multilingual Communication", "current_demand": "", "future_demand": ""},
    {"skill": "Cultural Intelligence & Diversity Awareness", "current_demand": "", "future_demand": ""},
    {"skill": "Privacy & Data Protection", "current_demand": "", "future_demand": ""},
    {"skill": "Legal Tech & RegTech Skills", "current_demand": "", "future_demand": ""},
    {"skill": "Creator Economy & Content Monetisation", "current_demand": "", "future_demand": ""},
    {"skill": "Autonomous Vehicles Technologies", "current_demand": "", "future_demand": ""},
    {"skill": "Robotics Engineering", "current_demand": "", "future_demand": ""},
    {"skill": "Quantum Cryptography", "current_demand": "", "future_demand": ""},
    {"skill": "Advanced Materials Science", "current_demand": "", "future_demand": ""},
    {"skill": "Biotechnology & Biomanufacturing", "current_demand": "", "future_demand": ""},
    {"skill": "Education Technology (EdTech)", "current_demand": "", "future_demand": ""},
    {"skill": "Spatial Computing / GIS", "current_demand": "", "future_demand": ""},
    {"skill": "Urban Planning & Smart Cities Technologies", "current_demand": "", "future_demand": ""},
    {"skill": "Autonomous Systems & Drones", "current_demand": "", "future_demand": ""},
    {"skill": "Quantum Computing Applications", "current_demand": "", "future_demand": ""},
]


@app.route("/industry_insights", methods=["GET", "POST"])
def industry_insights():
    filtered = []
    user_skill = ""
    if request.method == "POST":
        user_skill = request.form.get("skill", "").strip()
        filtered = [d for d in industry_data if d["skill"].lower() == user_skill.lower()]
    return render_template("industry_insights.html", filtered=filtered, user_skill=user_skill)

# ---------- LOGIN ----------
@app.route('/login', methods=['GET', 'POST'])
def login():
    if session.get('user'):
        return redirect(url_for('home'))

    error = None
    if request.method == 'POST':
        uname = request.form.get('username', '').strip().lower()
        pwd = request.form.get('password', '')
        user_record = users.get(uname)
        if user_record and user_record['password'] == pwd:
            session['user'] = {
                "name": uname.title(),
                "email": user_record['email'],
                "role": user_record['role']
            }
            return redirect(url_for('home'))
        error = "Invalid username or password"
    return render_template('login.html', error=error)


# ---------- REGISTER ----------
@app.route('/register', methods=['GET', 'POST'])
def register():
    if session.get('user'):
        return redirect(url_for('home'))

    error = None
    if request.method == 'POST':
        uname = request.form.get('username', '').strip().lower()
        email = request.form.get('email', '').strip()
        pwd = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if not uname or not email or not pwd:
            error = "All fields are required."
        elif pwd != confirm:
            error = "Passwords do not match."
        elif uname in users:
            error = "User already exists. Please log in."
        else:
            users[uname] = {"password": pwd, "email": email, "role": "New User"}
            session['user'] = {
                "name": uname.title(),
                "email": email,
                "role": "New User"
            }
            return redirect(url_for('home'))

    return render_template('register.html', error=error)


# ---------- CAREER FORM ----------
@app.route('/career_form', methods=['GET', 'POST'])
def career_form():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        skills_text = request.form.get('skills', '').strip()
        interests = request.form.get('interests', '').strip()
        resume = request.files.get('resume')

        extracted_text = ""

        # Resume upload
        if resume and allowed_file(resume.filename):
            filename = secure_filename(resume.filename)
            filepath = os.path.join('temp', filename)
            os.makedirs('temp', exist_ok=True)
            resume.save(filepath)

            if filename.endswith('.pdf'):
                extracted_text = extract_text_from_pdf(filepath)
            elif filename.endswith('.docx'):
                extracted_text = extract_text_from_docx(filepath)

            os.remove(filepath)

        # Combine all text
        combined_text = f"{skills_text} {interests} {extracted_text}"
        user_skills = [w.strip() for w in combined_text.split() if len(w) > 2]

        # Generate top 5 suggestions
        suggestions = suggest_career(user_skills)

        return render_template(
            'career_result.html',
            name=name,
            skills=user_skills,
            career_matches=suggestions
        )

    return render_template('career_form.html')



# ---------- LOGOUT ----------
@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('home'))


@app.route('/profile')
def profile():
    user = session.get("user")
    if not user:
        return redirect(url_for('login'))
    return render_template('profile.html', user=user)


if __name__ == '__main__':
    app.run(debug=True)
